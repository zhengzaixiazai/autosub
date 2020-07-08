#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Defines autosub's core functionality.
"""
# pylint: disable=too-many-lines
# Import built-in modules
import os
import multiprocessing
import time
import gettext
import gc
import re
import operator

# Import third-party modules
import progressbar
import pysubs2
import wcwidth
import docx
import googletrans
import auditok

# Any changes to the path and your own modules
from autosub import api_baidu
from autosub import api_google
from autosub import api_xfyun
from autosub import auditok_utils
from autosub import sub_utils
from autosub import ffmpeg_utils
from autosub import constants
from autosub import exceptions

CORE_TEXT = gettext.translation(domain=__name__,
                                localedir=constants.LOCALE_PATH,
                                languages=[constants.CURRENT_LOCALE],
                                fallback=True)

_ = CORE_TEXT.gettext


def auditok_opt_opt(  # pylint: disable=too-many-locals, too-many-branches, too-many-statements
        config_dict,
        audio_wav,
        concurrency=constants.DEFAULT_CONCURRENCY):
    """
    Function for optimize auditok options.
    """

    auditok_utils.validate_astats_config(config_dict)

    delta_et = (config_dict["max_et"] - config_dict["min_et"]) / (config_dict["et_pass"] + 1)
    delta_mxcs =\
        (config_dict["max_mxcs"] - config_dict["min_mxcs"]) / (config_dict["mxcs_pass"] + 1)
    input_stats = []
    asource = auditok.ADSFactory.ads(
        filename=audio_wav, record=True)

    et_i = config_dict["min_et"] + delta_et
    while et_i < config_dict["max_et"]:
        mxcs_i = config_dict["min_mxcs"] + delta_mxcs
        while mxcs_i < config_dict["max_mxcs"]:
            input_stats.append(auditok_utils.AuditokSTATS(
                energy_t=et_i,
                mxcs=mxcs_i,
                mnrs=config_dict["mnrs"],
                mxrs=config_dict["mxrs"],
                nsml=config_dict["nsml"],
                dts=config_dict["dts"],
                audio_wav=audio_wav
            ))
            mxcs_i = mxcs_i + delta_mxcs
        et_i = et_i + delta_et

    pool = multiprocessing.Pool(concurrency)
    widgets = [_("Auditok options optimization: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(input_stats)).start()

    try:
        i = 0
        tasks = []
        result_stats = []
        for stat in input_stats:
            tasks.append(pool.apply_async(
                auditok_utils.auditok_gen_stats_regions,
                args=(stat, asource)))
            gc.collect(0)

        for task in tasks:
            i = i + 1
            result_stats.append(task.get())
            pbar.update(i)

        rank_list = [
            sorted(result_stats, key=operator.attrgetter('big_region_count')),
            sorted(result_stats, key=operator.attrgetter('small_region_count')),
            sorted(result_stats, key=operator.attrgetter('delta_region_size'))]

        for stats_ in result_stats:
            for rank_item in rank_list:
                stats_.rank_count = rank_item.index(stats_) + stats_.rank_count

        result = min(result_stats)
        asource.close()
        pbar.finish()
        print(_("Best options for Auditok is:\n"
                "mxcs = {mxcs}s\net = {et}").format(mxcs=result.mxcs, et=result.et))
        config_dict["result_mxcs"] = result.mxcs
        config_dict["result_et"] = result.et
        pool.terminate()
        pool.join()
        return result.events

    except KeyboardInterrupt:
        asource.close()
        pbar.finish()
        pool.terminate()
        pool.join()
        return None


def trim_audio_regions(  # pylint: disable=too-many-arguments, too-many-locals, too-many-branches
        audio_fragments,
        events,
        delta,
        is_keep=False,
        trim_size=constants.DEFAULT_MIN_REGION_SIZE,
        energy_threshold=constants.DEFAULT_ENERGY_THRESHOLD,
        min_region_size=constants.DEFAULT_MIN_REGION_SIZE,
        max_region_size=constants.DEFAULT_MAX_REGION_SIZE,
        max_continuous_silence=constants.DEFAULT_CONTINUOUS_SILENCE,
        mode=auditok.StreamTokenizer.STRICT_MIN_LENGTH):
    """
    Give input audio fragments and trim the events.
    """

    widgets = [_("Trimming events: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(events)).start()
    try:
        i = 0
        regions = []
        for audio_fragment in audio_fragments:
            regions.append(auditok_utils.auditok_gen_speech_regions(
                audio_fragment,
                energy_threshold,
                min_region_size,
                max_region_size,
                max_continuous_silence,
                mode))
            gc.collect(0)

        for region in regions:
            if region:
                if events[i].start > delta:
                    start_delta = events[i].start - delta
                else:
                    start_delta = events[i].start
                start = 0
                end = 0
                if len(region) > 1:
                    if region[0][1] - region[0][0] <= trim_size:
                        start = start_delta + region[1][0]
                    if region[-1][1] - region[-1][0] <= trim_size:
                        end = start_delta + region[-2][1]
                if not end:
                    end = start_delta + region[-1][1]
                if not start:
                    start = start_delta + region[0][0]
                events[i].start = start
                events[i].end = end
            i = i + 1
            pbar.update(i)

        i = 0
        events_len = len(events)
        while i < events_len:
            if i > 0:
                if events[i].start < events[i - 1].end:
                    events[i].start = events[i - 1].end
            if i < events_len - 1:
                if events[i].end > events[i + 1].start:
                    events[i].end = events[i + 1].start
            i = i + 1

        if not is_keep:
            for audio_fragment in audio_fragments:
                os.remove(audio_fragment)

        pbar.finish()

    except KeyboardInterrupt:
        pbar.finish()


def bulk_audio_conversion(  # pylint: disable=too-many-arguments, too-many-locals
        source_file,
        regions,
        split_cmd,
        suffix,
        concurrency=constants.DEFAULT_CONCURRENCY,
        output=None,
        is_keep=False,
        include_before=0.0,
        include_after=0.0):
    """
    Give an input audio/video file and
    generate short-term audio fragments.
    """

    if not regions:
        return None

    pool = multiprocessing.Pool(concurrency)

    converter = ffmpeg_utils.SplitIntoAudioPiece(
        source_path=source_file,
        cmd=split_cmd,
        suffix=suffix,
        output=output,
        is_keep=is_keep,
        include_before=include_before,
        include_after=include_after)

    print(_("\nConverting speech regions to short-term fragments."))
    widgets = [_("Converting: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(regions)).start()
    try:
        audio_fragments = []
        for i, audio_fragment in enumerate(pool.imap(converter, regions)):
            if audio_fragment:
                audio_fragments.append(audio_fragment)
            pbar.update(i)
            gc.collect(0)
        pbar.finish()
        pool.terminate()
        pool.join()

    except KeyboardInterrupt:
        pbar.finish()
        pool.terminate()
        pool.join()
        return None
    return audio_fragments


def gsv2_to_text(  # pylint: disable=too-many-locals,too-many-arguments,too-many-branches,too-many-statements
        audio_fragments,
        api_url,
        headers,
        concurrency=constants.DEFAULT_CONCURRENCY,
        min_confidence=0.0,
        is_keep=False,
        result_list=None):
    """
    Give a list of short-term audio fragment files
    and generate text_list from Google speech-to-text V2 api.
    """
    text_list = []
    pool = multiprocessing.Pool(concurrency)

    recognizer = api_google.GoogleSpeechV2(
        api_url=api_url,
        headers=headers,
        min_confidence=min_confidence,
        is_keep=is_keep,
        is_full_result=result_list is not None)

    print(_("\nSending short-term fragments to Google Speech V2 API and getting result."))
    widgets = [_("Speech-to-Text: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(audio_fragments)).start()
    try:
        # get transcript
        if result_list is None:
            for i, transcript in enumerate(pool.imap(recognizer, audio_fragments)):
                if transcript:
                    text_list.append(transcript)
                else:
                    text_list.append("")
                gc.collect(0)
                pbar.update(i)
        # get full result and transcript
        else:
            for i, result in enumerate(pool.imap(recognizer, audio_fragments)):
                if result:
                    result_list.append(result)
                    transcript = \
                        api_google.get_google_speech_v2_transcript(min_confidence, result)
                    if transcript:
                        text_list.append(transcript)
                        continue
                else:
                    result_list.append("")
                text_list.append("")
                gc.collect(0)
                pbar.update(i)
        pbar.finish()
        pool.terminate()
        pool.join()

    except (KeyboardInterrupt, AttributeError) as error:
        pbar.finish()
        pool.terminate()
        pool.join()

        if error == AttributeError:
            print(
                _("Error: Connection error happened too many times.\nAll works done."))

        return None

    return text_list


def gcsv1_to_text(  # pylint: disable=too-many-locals,too-many-arguments,too-many-branches,too-many-statements, too-many-nested-blocks
        audio_fragments,
        sample_rate,
        api_url=None,
        headers=None,
        config=None,
        concurrency=constants.DEFAULT_CONCURRENCY,
        src_language=constants.DEFAULT_SRC_LANGUAGE,
        min_confidence=0.0,
        is_keep=False,
        result_list=None):
    """
    Give a list of short-term audio fragment files
    and generate text_list from Google cloud speech-to-text V1P1Beta1 api.
    """

    text_list = []
    pool = multiprocessing.Pool(concurrency)

    print(_("\nSending short-term fragments to Google Cloud Speech V1P1Beta1 API"
            " and getting result."))
    widgets = [_("Speech-to-Text: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(audio_fragments)).start()

    try:
        if api_url:
            # https://cloud.google.com/speech-to-text/docs/reference/rest/v1p1beta1/RecognitionConfig
            if config:
                # Use the fixed argument
                if "languageCode" in config:
                    config["languageCode"] = src_language
                else:
                    config["language_code"] = src_language
            else:
                config = {
                    "encoding": api_google.google_ext_to_enc(audio_fragments[0]),
                    "sampleRateHertz": sample_rate,
                    "languageCode": src_language}

            recognizer = api_google.GCSV1P1Beta1URL(
                config=config,
                api_url=api_url,
                headers=headers,
                min_confidence=min_confidence,
                is_keep=is_keep,
                is_full_result=result_list is not None)

            # get transcript
            if result_list is None:
                for i, transcript in enumerate(pool.imap(recognizer, audio_fragments)):
                    if transcript:
                        text_list.append(transcript)
                    else:
                        text_list.append("")
                    pbar.update(i)
            # get full result and transcript
            else:
                for i, result in enumerate(pool.imap(recognizer, audio_fragments)):
                    if result:
                        result_list.append(result)
                        transcript = api_google.get_gcsv1p1beta1_transcript(
                            min_confidence,
                            result)
                        if transcript:
                            text_list.append(transcript)
                            continue
                    else:
                        result_list.append("")
                    text_list.append("")
                    pbar.update(i)

        else:
            # https://googleapis.dev/python/speech/latest/gapic/v1/types.html#google.cloud.speech_v1.types.RecognitionConfig
            if config:
                # Use the fixed arguments
                config["encoding"] = api_google.google_ext_to_enc(
                    extension=audio_fragments[0],
                    is_string=False
                )
                config["language_code"] = src_language
            else:
                config = {
                    "encoding": api_google.google_ext_to_enc(
                        extension=audio_fragments[0],
                        is_string=False),
                    "sample_rate_hertz": sample_rate,
                    "language_code": src_language}

            i = 0
            tasks = []
            for filename in audio_fragments:
                # google cloud speech-to-text client can't use multiprocessing.pool
                # based on class call, otherwise will receive pickling error
                tasks.append(pool.apply_async(
                    api_google.gcsv1p1beta1_service_client,
                    args=(filename, is_keep, config, min_confidence,
                          result_list is not None)))
                gc.collect(0)

            if result_list is None:
                for task in tasks:
                    i = i + 1
                    transcript = task.get()
                    if transcript:
                        text_list.append(transcript)
                    else:
                        text_list.append("")
                    pbar.update(i)
            else:
                for task in tasks:
                    i = i + 1
                    result = task.get()
                    result_list.append(result)
                    transcript = api_google.get_gcsv1p1beta1_transcript(
                        min_confidence,
                        result)
                    if transcript:
                        text_list.append(transcript)
                    else:
                        text_list.append("")
                    pbar.update(i)

        pbar.finish()
        pool.terminate()
        pool.join()

    except (KeyboardInterrupt, AttributeError) as error:
        pbar.finish()
        pool.terminate()
        pool.join()

        if error == AttributeError:
            print(
                _("Error: Connection error happened too many times.\nAll works done."))

        return None

    except exceptions.SpeechToTextException as err_msg:
        pbar.finish()
        pool.terminate()
        pool.join()
        print(_("Receive something unexpected:"))
        print(err_msg)
        return None

    return text_list


def xfyun_to_text(  # pylint: disable=too-many-locals, too-many-arguments,
        # pylint: disable=too-many-branches, too-many-statements, too-many-nested-blocks
        audio_fragments,
        config,
        concurrency=constants.DEFAULT_CONCURRENCY,
        is_keep=False,
        result_list=None):
    """
    Give a list of short-term audio fragment files
    and generate text_list from Google cloud speech-to-text V1P1Beta1 api.
    """

    text_list = []

    if "api_address" in config:
        api_address = config["api_address"]
    else:
        api_address = constants.XFYUN_SPEECH_WEBAPI_URL

    if "delete_chars" in config:
        delete_chars = config["delete_chars"]
    else:
        delete_chars = None

    pool = multiprocessing.Pool(concurrency)

    print(_("\nSending short-term fragments to Xun Fei Yun WebSocket API"
            " and getting result."))
    widgets = [_("Speech-to-Text: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(audio_fragments)).start()

    try:
        recognizer = api_xfyun.XfyunWebSocketAPI(
            app_id=config["app_id"],
            api_key=config["api_key"],
            api_secret=config["api_secret"],
            api_address=api_address,
            business_args=config["business"],
            is_full_result=result_list is not None,
            delete_chars=delete_chars)

        # get transcript
        if result_list is None:
            for i, transcript in enumerate(pool.imap(recognizer, audio_fragments)):
                if transcript:
                    text_list.append(transcript)
                else:
                    text_list.append("")
                pbar.update(i)
        # get full result and transcript
        else:
            for i, result in enumerate(pool.imap(recognizer, audio_fragments)):
                if result:
                    result_list.append(result)
                    transcript = ""
                    for item in result:
                        transcript = transcript + api_xfyun.get_xfyun_transcript(
                            result_dict=item,
                            delete_chars=delete_chars)
                    if transcript:
                        text_list.append(transcript)
                        pbar.update(i)
                        continue
                else:
                    result_list.append("")
                text_list.append("")
                pbar.update(i)

        if not is_keep:
            for audio_fragment in audio_fragments:
                os.remove(audio_fragment)

        pbar.finish()
        pool.terminate()
        pool.join()

    except (KeyboardInterrupt, AttributeError) as error:
        if not is_keep:
            for audio_fragment in audio_fragments:
                os.remove(audio_fragment)
        pbar.finish()
        pool.terminate()
        pool.join()

        if error == AttributeError:
            print(
                _("Error: Connection error happened too many times.\nAll works done."))

        return None

    except exceptions.SpeechToTextException as err_msg:
        if not is_keep:
            for audio_fragment in audio_fragments:
                os.remove(audio_fragment)
        pbar.finish()
        pool.terminate()
        pool.join()
        print(_("Receive something unexpected:"))
        print(err_msg)
        return None

    return text_list


def baidu_to_text(  # pylint: disable=too-many-locals, too-many-arguments,
        # pylint: disable=too-many-branches, too-many-statements, too-many-nested-blocks
        audio_fragments,
        config,
        concurrency=constants.DEFAULT_CONCURRENCY,
        is_keep=False,
        result_list=None):
    """
    Give a list of short-term audio fragment files
    and generate text_list from Google cloud speech-to-text V1P1Beta1 api.
    """

    text_list = []

    if config["config"]["dev_pid"] == 80001:
        # pro edition of baidu asr
        api_url = constants.BAIDU_PRO_ASR_URL
        print(_("\nSending short-term fragments to Baidu PRO ASR API"
                " and getting result."))
    else:
        print(_("\nSending short-term fragments to Baidu ASR API"
                " and getting result."))
        api_url = constants.BAIDU_ASR_URL

    if "delete_chars" in config:
        delete_chars = config["delete_chars"]
    else:
        delete_chars = None

    try:
        if "token" not in config["config"]:
            print(_("Get the token online."))
            config["config"]["token"] = \
                api_baidu.get_baidu_token(api_secret=config["api_secret"],
                                          api_key=config["api_key"])
        else:
            print(_("Use the token from the config."))

    except exceptions.SpeechToTextException as err_msg:
        print(_("Failed to get the token. Error message:"))
        print(err_msg)
        return None

    pool = multiprocessing.Pool(concurrency)

    widgets = [_("Speech-to-Text: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=len(audio_fragments)).start()

    try:
        recognizer = api_baidu.BaiduASRAPI(
            config=config["config"],
            api_url=api_url,
            is_keep=is_keep,
            is_full_result=result_list is not None,
            delete_chars=delete_chars)

        # get transcript
        if result_list is None:
            for i, transcript in enumerate(pool.imap(recognizer, audio_fragments)):
                if transcript:
                    text_list.append(transcript)
                else:
                    text_list.append("")
                pbar.update(i)
        # get full result and transcript
        else:
            for i, result in enumerate(pool.imap(recognizer, audio_fragments)):
                if result:
                    result_list.append(result)
                    transcript = api_baidu.get_baidu_transcript(
                        result, delete_chars)
                    if transcript:
                        text_list.append(transcript)
                        pbar.update(i)
                        continue
                else:
                    result_list.append("")
                text_list.append("")
                pbar.update(i)
        pbar.finish()
        pool.terminate()
        pool.join()

    except (KeyboardInterrupt, AttributeError) as error:
        pbar.finish()
        pool.terminate()
        pool.join()

        if error == AttributeError:
            print(
                _("Error: Connection error happened too many times.\nAll works done."))

        return None

    except exceptions.SpeechToTextException as err_msg:
        pbar.finish()
        pool.terminate()
        pool.join()
        print(_("Receive something unexpected:"))
        print(err_msg)
        return None

    return text_list


def list_to_googletrans(  # pylint: disable=too-many-locals, too-many-arguments, too-many-branches, too-many-statements
        text_list,
        translator,
        src_language=constants.DEFAULT_SRC_LANGUAGE,
        dst_language=constants.DEFAULT_DST_LANGUAGE,
        size_per_trans=constants.DEFAULT_SIZE_PER_TRANS,
        sleep_seconds=constants.DEFAULT_SLEEP_SECONDS,
        drop_override_codes=False,
        delete_chars=None):
    """
    Give a text list, generate translated text list from GoogleTranslatorV2 api.
    """

    if not text_list:
        return None

    size = 0
    i = 0
    partial_index = []
    valid_index = []
    is_last = ""
    text_list_length = len(text_list)
    if size_per_trans <= 0:
        size_per_trans = float("inf")
    while i < text_list_length:
        if text_list[i]:
            if not is_last:
                is_last = text_list[i]
                valid_index.append(i)
                # valid_index for valid text position start
            wcswidth_text = wcwidth.wcswidth(text_list[i])
            if wcswidth_text * 10 / len(text_list[i]) >= 10:
                # If text contains full-wide char,
                # count its length about 4 times than the ordinary text.
                # Avoid weird problem when text has full-wide char.
                # In this case Google will count a full-wide char
                # at least 2 times larger than a half-wide char.
                # It will certainly exceed the limit of the size_per_trans.
                # Causing a googletrans internal jsondecode error.
                size = size + wcswidth_text * 2
            else:
                size = size + len(text_list[i])
            if size > size_per_trans:
                # use size_per_trans to split the list
                partial_index.append(i)
                size = 0
                continue
                # stay at this line of text
                # in case it's the last one
        else:
            if is_last:
                is_last = text_list[i]
                valid_index.append(i)
                # valid_index for valid text position end
        i = i + 1

    if size:
        partial_index.append(i)
        # python sequence
        # every group's end index
    else:
        return None

    len_valid_index = len(valid_index)

    if len_valid_index % 2:
        valid_index.append(i)
        # valid_index for valid text position end

    if translator != ManualTranslator and src_language == "auto":
        content_to_trans = '\n'.join(text_list[i:partial_index[0]])
        result_src = translator.detect(content_to_trans).lang
    else:
        result_src = src_language

    print(_("\nTranslating text from \"{0}\" to \"{1}\".").format(
        result_src,
        dst_language))

    widgets = [_("Translation: "),
               progressbar.Percentage(), ' ',
               progressbar.Bar(), ' ',
               progressbar.ETA()]
    pbar = progressbar.ProgressBar(widgets=widgets, maxval=i).start()

    try:
        translated_text = []
        i = 0
        # total position
        j = 0
        # valid_index position

        for index in partial_index:
            content_to_trans = '\n'.join(text_list[i:index])
            if drop_override_codes:
                content_to_trans = "".join(re.compile(r'{.*?}').split(content_to_trans))
            translation = translator.translate(text=content_to_trans,
                                               dest=dst_language,
                                               src=src_language)
            result_text = translation.text.translate(str.maketrans('’', '\''))
            result_src = translation.src
            result_list = result_text.split('\n')
            k = 0
            len_result_list = len(result_list)
            while i < index and j < len_valid_index and k < len_result_list:
                if not result_list[k]:
                    # if the result is invalid,
                    # continue
                    k = k + 1
                    continue
                if i < valid_index[j]:
                    # if text is invalid,
                    # append the empty string
                    # and then continue
                    translated_text.append("")
                    i = i + 1
                    if i % 20 == 5:
                        pbar.update(i)
                    continue
                if i < valid_index[j + 1]:
                    # if text is valid, append it
                    if delete_chars:
                        result_list[k] = result_list[k].translate(
                            str.maketrans(delete_chars, " " * len(delete_chars)))
                        result_list[k] = result_list[k].rstrip(" ")
                    translated_text.append(result_list[k])
                    k = k + 1
                    i = i + 1
                else:
                    j = j + 2
                if i % 20 == 5:
                    pbar.update(i)
            if len(partial_index) > 1:
                time.sleep(sleep_seconds)

        i = valid_index[-1]
        while i < partial_index[-1]:
            # if valid_index[-1] is less than partial_index[-1]
            # add empty strings
            translated_text.append("")
            i = i + 1

        pbar.finish()

    except KeyboardInterrupt:
        pbar.finish()
        print(_("Cancelling translation."))
        return 1

    return translated_text, result_src


class ManualTranslator:  # pylint: disable=too-few-public-methods
    """
    Class for performing translation manually.
    """
    def __init__(self,
                 trans_doc_name,
                 input_m=input):
        # pylint: disable=too-many-arguments
        self.trans_doc_name = trans_doc_name
        self.input_m = input_m

    def translate(self,
                  text,
                  dest=None,
                  src=None):
        """
        Translate text manually.
        """
        if self.trans_doc_name.endswith("docx"):
            trans_doc = docx.Document()
            trans_doc.add_paragraph(text=text)
            trans_doc.save(self.trans_doc_name)
            if self.input_m:
                self.input_m(_("Wait for the manual translation. "
                               "Press Enter to continue."))
            else:
                print(_("Wait 20 seconds for the manual translation. "))
                widgets = [_("Manual translation: "),
                           progressbar.Percentage(), ' ',
                           progressbar.Bar(), ' ',
                           progressbar.ETA()]
                pbar = progressbar.ProgressBar(widgets=widgets, maxval=20).start()
                for i in range(20):
                    pbar.update(i)
                    time.sleep(1)
                pbar.finish()
            trans_doc = docx.Document(self.trans_doc_name)
            para_list = []
            for para in trans_doc.paragraphs:
                para_list.append(para.text)
            trans_doc_str = '\n'.join(para_list)
        else:
            trans_doc_name = sub_utils.str_to_file(
                str_=text,
                output=self.trans_doc_name,
                input_m=self.input_m)
            if self.input_m:
                self.input_m(_("Wait for the manual translation. "
                               "Press Enter to continue."))
            else:
                print(_("Wait 20 seconds for the manual translation. "))
                widgets = [_("Manual translation: "),
                           progressbar.Percentage(), ' ',
                           progressbar.Bar(), ' ',
                           progressbar.ETA()]
                pbar = progressbar.ProgressBar(widgets=widgets, maxval=20).start()
                for i in range(20):
                    pbar.update(i)
                    time.sleep(1)
                pbar.finish()
            trans_doc = open(trans_doc_name, encoding=constants.DEFAULT_ENCODING)
            trans_doc_str = trans_doc.read()
            trans_doc.close()
        constants.delete_path(self.trans_doc_name)
        return googletrans.client.Translated(
            src=src, dest=dest, origin="manual",
            text=trans_doc_str, pronunciation="manual", extra_data="manual")


def list_to_sub_str(
        timed_text,
        fps=30.0,
        subtitles_file_format=constants.DEFAULT_SUBTITLES_FORMAT):
    """
    Give an input timed text list, format it to a string.
    """

    if subtitles_file_format in ('srt', 'tmp', 'ass', 'ssa'):
        pysubs2_obj = pysubs2.SSAFile()
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=timed_text)
        formatted_subtitles = pysubs2_obj.to_string(
            format_=subtitles_file_format)

    elif subtitles_file_format == 'vtt':
        formatted_subtitles = sub_utils.list_to_vtt_str(
            subtitles=timed_text)

    elif subtitles_file_format == 'json':
        formatted_subtitles = sub_utils.list_to_json_str(
            subtitles=timed_text)

    elif subtitles_file_format == 'ass.json':
        pysubs2_obj = pysubs2.SSAFile()
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=timed_text)
        formatted_subtitles = pysubs2_obj.to_string(
            format_='json')

    elif subtitles_file_format == 'txt':
        formatted_subtitles = sub_utils.list_to_txt_str(
            subtitles=timed_text)

    elif subtitles_file_format == 'sub':
        pysubs2_obj = pysubs2.SSAFile()
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=timed_text)
        formatted_subtitles = pysubs2_obj.to_string(
            format_='microdvd',
            fps=fps)
        # sub format need fps
        # ref https://pysubs2.readthedocs.io/en/latest
        # /api-reference.html#supported-input-output-formats

    elif subtitles_file_format == 'mpl2.txt':
        pysubs2_obj = pysubs2.SSAFile()
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=timed_text)
        formatted_subtitles = pysubs2_obj.to_string(
            format_='mpl2',
            fps=fps)

    else:
        # fallback process
        print(_("Format \"{fmt}\" not supported. "
                "Use \"{default_fmt}\" instead.").format(
                    fmt=subtitles_file_format,
                    default_fmt=constants.DEFAULT_SUBTITLES_FORMAT))
        pysubs2_obj = pysubs2.SSAFile()
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=timed_text)
        formatted_subtitles = pysubs2_obj.to_string(
            format_=constants.DEFAULT_SUBTITLES_FORMAT)

    return formatted_subtitles


def ssafile_to_sub_str(
        ssafile,
        fps=30.0,
        subtitles_file_format=constants.DEFAULT_SUBTITLES_FORMAT):
    """
    Give an input SSAFile, format it to a string.
    """

    if subtitles_file_format in ('srt', 'tmp', 'ass', 'ssa'):
        formatted_subtitles = ssafile.to_string(
            format_=subtitles_file_format)

    elif subtitles_file_format == 'vtt':
        formatted_subtitles = sub_utils.assfile_to_vtt_str(
            subtitles=ssafile)

    elif subtitles_file_format == 'json':
        formatted_subtitles = sub_utils.assfile_to_json_str(
            subtitles=ssafile)

    elif subtitles_file_format == 'ass.json':
        formatted_subtitles = ssafile.to_string(
            format_='json')

    elif subtitles_file_format == 'txt':
        formatted_subtitles = sub_utils.assfile_to_txt_str(
            subtitles=ssafile)

    elif subtitles_file_format == 'sub':
        formatted_subtitles = ssafile.to_string(
            format_='microdvd',
            fps=fps)
        # sub format need fps
        # ref https://pysubs2.readthedocs.io/en/latest
        # /api-reference.html#supported-input-output-formats

    elif subtitles_file_format == 'mpl2.txt':
        formatted_subtitles = ssafile.to_string(
            format_='mpl2',
            fps=fps)

    else:
        # fallback process
        print(_("Format \"{fmt}\" not supported. "
                "Use \"{default_fmt}\" instead.").format(
                    fmt=subtitles_file_format,
                    default_fmt=constants.DEFAULT_SUBTITLES_FORMAT))
        formatted_subtitles = ssafile.to_string(
            format_=constants.DEFAULT_SUBTITLES_FORMAT)

    return formatted_subtitles


def list_to_ass_str(
        text_list,
        styles_list,
        subtitles_file_format=constants.DEFAULT_SUBTITLES_FORMAT,
        same_event_type=0):
    """
    Give an input timed text list, format it to an ass string.
    """
    pysubs2_obj = pysubs2.SSAFile()
    pysubs2_obj.styles = \
        {styles_list[i]: styles_list[i + 1] for i in range(0, len(styles_list), 2)}
    if not isinstance(text_list[0], list):
        # text_list is [((start, end), text), ...]
        # text_list provides regions
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=text_list,
            style_name=styles_list[0])
    else:
        # text_list is [[src_list], [dst_list]]
        # src_list provides regions
        sub_utils.pysubs2_ssa_event_add(
            src_ssafile=None,
            dst_ssafile=pysubs2_obj,
            text_list=text_list[0],
            style_name=styles_list[0])
        src_obj = pysubs2_obj
        pysubs2_obj = pysubs2.SSAFile()
        if len(styles_list) == 1:
            sub_utils.pysubs2_ssa_event_add(
                src_ssafile=src_obj,
                dst_ssafile=pysubs2_obj,
                text_list=text_list[1],
                style_name=styles_list[0],
                same_event_type=same_event_type)
        else:
            sub_utils.pysubs2_ssa_event_add(
                src_ssafile=src_obj,
                dst_ssafile=pysubs2_obj,
                text_list=text_list[1],
                style_name=styles_list[2],
                same_event_type=same_event_type)

    if subtitles_file_format != 'ass.json':
        formatted_subtitles = pysubs2_obj.to_string(format_=subtitles_file_format)
    else:
        formatted_subtitles = pysubs2_obj.to_string(format_='json')

    return formatted_subtitles
